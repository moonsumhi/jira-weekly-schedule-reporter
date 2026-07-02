"""Document management — folder/file upload, listing, full-text search."""

import io
import logging
import mimetypes
import os
import re
import shutil
import subprocess
import tempfile
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import List, Optional

from urllib.parse import quote

from bson import ObjectId
from fastapi import APIRouter, Depends, File, Form, HTTPException, Query, Request, UploadFile
from fastapi.responses import FileResponse, HTMLResponse

from app.db.mongo import MongoClientManager
from app.models.user import UserPublic
from app.routers.auth import get_current_user, get_user_by_email
from app.core.security import decode_token
from app.utils.mongo import fmt_dt
from app.utils.mongo import oid as parse_oid

logger = logging.getLogger(__name__)
router = APIRouter()

UPLOAD_DIR = Path("/app/uploads/documents")
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)


def _now() -> datetime:
    return datetime.now(timezone.utc)


def _db():
    return MongoClientManager.get_db()


# ── Text extraction ────────────────────────────────────────────────────────────

def _extract_text(content: bytes, extension: str) -> str:
    ext = extension.lower().lstrip(".")
    try:
        if ext == "pdf":
            return _extract_pdf(content)
        elif ext in ("xlsx", "xls"):
            return _extract_excel(content)
        elif ext == "hwp":
            return _extract_hwp(content)
        elif ext == "docx":
            return _extract_docx(content)
        elif ext in ("txt", "md", "csv"):
            return content.decode("utf-8", errors="ignore")
    except Exception as e:
        logger.warning("Text extraction failed for .%s: %s", ext, e)
    return ""


def _extract_pdf(content: bytes) -> str:
    import fitz
    doc = fitz.open(stream=content, filetype="pdf")
    return "\n".join(page.get_text() for page in doc)


def _extract_excel(content: bytes) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    parts = []
    for ws in wb.worksheets:
        parts.append(f"[{ws.title}]")
        for row in ws.iter_rows(values_only=True):
            row_text = " ".join(str(c) for c in row if c is not None)
            if row_text.strip():
                parts.append(row_text)
    return "\n".join(parts)


def _extract_hwp(content: bytes) -> str:
    with tempfile.NamedTemporaryFile(suffix=".hwp", delete=False) as f:
        f.write(content)
        tmp_path = f.name
    try:
        result = subprocess.run(
            ["hwp5txt", tmp_path],
            capture_output=True, text=True, timeout=30,
        )
        return result.stdout
    except Exception as e:
        logger.warning("HWP extraction failed: %s", e)
        return ""
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass


def _extract_docx(content: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(content))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


# ── Folder helpers ─────────────────────────────────────────────────────────────

async def _ensure_folder_path(db, path_parts: list, user_email: str, parent_id: Optional[str] = None) -> Optional[str]:
    """Ensure folder hierarchy exists, return leaf folder id string.
    parent_id: starting parent folder (None = root). Used to scope uploads under a specific folder.
    """
    for part in path_parts:
        existing = await db["document_folders"].find_one(
            {"name": part, "parent_id": parent_id}
        )
        if existing:
            parent_id = str(existing["_id"])
        else:
            doc = {
                "name": part,
                "parent_id": parent_id,
                "created_at": _now(),
                "created_by": user_email,
            }
            result = await db["document_folders"].insert_one(doc)
            parent_id = str(result.inserted_id)
    return parent_id


# ── Serializer ────────────────────────────────────────────────────────────────

def _file_out(f: dict, include_text: bool = False) -> dict:
    out = {
        "id": str(f["_id"]),
        "name": f["name"],
        "folder_id": f.get("folder_id"),
        "extension": f.get("extension", ""),
        "mime_type": f.get("mime_type", ""),
        "size": f.get("size", 0),
        "created_at": fmt_dt(f.get("created_at")),
        "created_by": f.get("created_by"),
        "converted_from": f.get("converted_from"),
    }
    if include_text:
        out["text_content"] = f.get("text_content", "")
    return out


def _lo_convert(src: Path, target_format: str, out_dir: Path) -> Path:
    """LibreOffice headless로 파일 변환.
    HWP 변환 순서:
      1) LibreOffice 직접 (LO 25+ 네이티브 HWP 필터)
      2) hwp5html → HTML → target (표 구조 보존 우수)
      3) hwp5odt → ODT → target (최후 수단)
    """
    lo_home = Path(tempfile.mkdtemp(prefix="lo_home_"))
    try:
        actual_src = src
        if src.suffix.lower() == ".hwp":
            # 1) LibreOffice 직접 변환
            lo_direct = subprocess.run(
                ["libreoffice", "--headless", "--convert-to", target_format,
                 "--outdir", str(out_dir), str(src)],
                capture_output=True, text=True, timeout=120,
                env={**os.environ, "HOME": str(lo_home)},
            )
            direct_out = out_dir / f"{src.stem}.{target_format}"
            if lo_direct.returncode == 0 and direct_out.exists():
                return direct_out

            # 2) hwp5html → HTML → target (표 보존이 ODT 경로보다 우수)
            html_path = out_dir / f"{src.stem}.html"
            r_html = subprocess.run(
                ["hwp5html", "--output", str(html_path), str(src)],
                capture_output=True, text=True, timeout=120,
            )
            if r_html.returncode == 0 and html_path.exists():
                lo_html = subprocess.run(
                    ["libreoffice", "--headless", "--convert-to", target_format,
                     "--outdir", str(out_dir), str(html_path)],
                    capture_output=True, text=True, timeout=120,
                    env={**os.environ, "HOME": str(lo_home)},
                )
                html_out = out_dir / f"{src.stem}.{target_format}"
                if lo_html.returncode == 0 and html_out.exists():
                    return html_out

            # 3) hwp5odt → ODT → target (RelaxNG 경고 무시, 파일 존재 여부만 확인)
            odt_path = out_dir / f"{src.stem}.odt"
            subprocess.run(
                ["hwp5odt", "--output", str(odt_path), str(src)],
                capture_output=True, text=True, timeout=120,
            )
            if not odt_path.exists():
                raise RuntimeError("HWP 변환 실패: 모든 변환 방법이 실패했습니다.")
            actual_src = odt_path

        result = subprocess.run(
            ["libreoffice", "--headless", "--convert-to", target_format,
             "--outdir", str(out_dir), str(actual_src)],
            capture_output=True, text=True, timeout=120,
            env={**os.environ, "HOME": str(lo_home)},
        )
        if result.returncode != 0:
            raise RuntimeError(result.stderr or result.stdout)
    finally:
        shutil.rmtree(lo_home, ignore_errors=True)
    out_file = out_dir / f"{actual_src.stem}.{target_format}"
    if not out_file.exists():
        raise RuntimeError(f"변환된 파일을 찾을 수 없습니다: {out_file}")
    return out_file


def _snippet(text: str, q: str, context: int = 120) -> str:
    idx = text.lower().find(q.lower())
    if idx == -1:
        return text[:context * 2] if text else ""
    start = max(0, idx - context)
    end = min(len(text), idx + len(q) + context)
    s = text[start:end]
    if start > 0:
        s = "…" + s
    if end < len(text):
        s = s + "…"
    return s


# ── Endpoints ─────────────────────────────────────────────────────────────────

@router.post("/upload")
async def upload_files(
    files: List[UploadFile] = File(...),
    paths: List[str] = Form(...),
    target_folder_id: Optional[str] = Form(None),
    current_user: UserPublic = Depends(get_current_user),
):
    """Upload files preserving folder structure (paths = relative paths per file).
    target_folder_id: if provided, files with no folder in their path are placed here.
    """
    db = _db()
    uploaded = []

    for file, rel_path in zip(files, paths):
        content = await file.read()
        parts = Path(rel_path).parts  # e.g. ("폴더", "서브폴더", "파일.pdf")

        folder_id: Optional[str] = None
        if len(parts) > 1:
            folder_id = await _ensure_folder_path(
                db, list(parts[:-1]), current_user.email,
                parent_id=target_folder_id  # scoped 모드: 스코프 루트 하위에 생성
            )
        elif target_folder_id:
            folder_id = target_folder_id

        filename = parts[-1]
        extension = Path(filename).suffix
        file_uuid = str(uuid.uuid4())
        save_path = UPLOAD_DIR / f"{file_uuid}{extension}"
        save_path.write_bytes(content)

        text_content = _extract_text(content, extension)
        mime_type = mimetypes.guess_type(filename)[0] or "application/octet-stream"

        doc = {
            "name": filename,
            "folder_id": folder_id,
            "extension": extension.lstrip(".").lower(),
            "mime_type": mime_type,
            "file_path": str(save_path),
            "text_content": text_content,
            "size": len(content),
            "is_deleted": False,
            "created_at": _now(),
            "created_by": current_user.email,
        }
        result = await db["document_files"].insert_one(doc)
        uploaded.append(str(result.inserted_id))

    return {"uploaded": len(uploaded), "ids": uploaded}


@router.post("/folders")
async def create_folder(
    name: str = Form(...),
    parent_id: Optional[str] = Form(None),
    current_user: UserPublic = Depends(get_current_user),
):
    db = _db()
    doc = {
        "name": name.strip(),
        "parent_id": parent_id or None,
        "created_at": _now(),
        "created_by": current_user.email,
    }
    result = await db["document_folders"].insert_one(doc)
    return {
        "id": str(result.inserted_id),
        "name": doc["name"],
        "parent_id": doc["parent_id"],
        "created_at": fmt_dt(doc["created_at"]),
    }


@router.get("/folders")
async def list_folders(current_user: UserPublic = Depends(get_current_user)):
    db = _db()
    folders = await db["document_folders"].find().to_list(length=None)
    return [
        {
            "id": str(f["_id"]),
            "name": f["name"],
            "parent_id": f.get("parent_id"),
            "created_at": fmt_dt(f.get("created_at")),
        }
        for f in folders
    ]


@router.get("/root-files")
async def list_root_files(current_user: UserPublic = Depends(get_current_user)):
    db = _db()
    files = await db["document_files"].find(
        {"folder_id": None, "is_deleted": {"$ne": True}}
    ).to_list(length=None)
    return [_file_out(f) for f in files]


@router.get("/folders/{folder_id}/files")
async def list_files_in_folder(
    folder_id: str,
    current_user: UserPublic = Depends(get_current_user),
):
    db = _db()
    files = await db["document_files"].find(
        {"folder_id": folder_id, "is_deleted": {"$ne": True}}
    ).to_list(length=None)
    return [_file_out(f) for f in files]


@router.get("/search")
async def search_documents(
    q: str = Query(..., min_length=1),
    current_user: UserPublic = Depends(get_current_user),
):
    db = _db()
    pattern = re.compile(re.escape(q), re.IGNORECASE)
    files = await db["document_files"].find(
        {
            "is_deleted": {"$ne": True},
            "$or": [
                {"name": pattern},
                {"text_content": pattern},
            ],
        }
    ).to_list(length=200)

    results = []
    for f in files:
        item = _file_out(f)
        item["snippet"] = _snippet(f.get("text_content", ""), q)
        results.append(item)
    return results


@router.get("/files/{file_id}")
async def get_file_meta(
    file_id: str,
    current_user: UserPublic = Depends(get_current_user),
):
    db = _db()
    f = await db["document_files"].find_one({"_id": parse_oid(file_id)})
    if not f:
        raise HTTPException(status_code=404, detail="파일을 찾을 수 없습니다.")
    return _file_out(f, include_text=True)


async def _get_user_any_auth(
    request: Request,
    token: Optional[str] = Query(None),
) -> UserPublic:
    """헤더 또는 쿼리 파라미터로 JWT 인증 (iframe PDF 뷰어용)."""
    # 쿼리 파라미터 토큰 우선, 없으면 Authorization 헤더 사용
    raw = token or request.headers.get("authorization", "").removeprefix("Bearer ").strip()
    try:
        email = decode_token(raw) if raw else None
    except Exception:
        email = None
    if not email:
        raise HTTPException(status_code=401, detail="Could not validate credentials")
    user = await get_user_by_email(email)
    if not user:
        raise HTTPException(status_code=401, detail="User not found")
    return UserPublic(
        id=str(user["_id"]),
        email=user["email"],
        full_name=user.get("full_name"),
        is_admin=bool(user.get("is_admin", False)),
        permissions=user.get("permissions", []),
        is_internal=False,
    )


@router.get("/files/{file_id}/content")
async def get_file_content(
    file_id: str,
    current_user: UserPublic = Depends(_get_user_any_auth),
):
    db = _db()
    f = await db["document_files"].find_one({"_id": parse_oid(file_id)})
    if not f:
        raise HTTPException(status_code=404, detail="파일을 찾을 수 없습니다.")

    file_path = Path(f["file_path"])
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="파일이 서버에 없습니다.")

    encoded_name = quote(f["name"])
    return FileResponse(
        path=str(file_path),
        media_type=f.get("mime_type", "application/octet-stream"),
        headers={"Content-Disposition": f"inline; filename*=UTF-8''{encoded_name}"},
    )


_HWP_BASE_CSS = """
<style>
* { box-sizing: border-box; }
body { font-family: 'Malgun Gothic', 'Noto Sans KR', sans-serif; font-size: 10pt;
       line-height: 1.6; padding: 24px; color: #222; }
table { border-collapse: collapse; width: auto; margin: 8px 0; }
td, th { border: 1px solid #888; padding: 4px 8px; vertical-align: top; }
th { background: #e8eaf6; font-weight: bold; }
p { margin: 2px 0; }
img { max-width: 100%; height: auto; }
</style>
"""

def _make_self_contained(html: str, base_dir: Path) -> str:
    """CSS 인라인 삽입 + 상대 경로 이미지를 base64 data URL로 치환."""
    import base64 as _b64

    mime_map = {
        'png': 'image/png', 'jpg': 'image/jpeg', 'jpeg': 'image/jpeg',
        'gif': 'image/gif', 'bmp': 'image/bmp', 'svg': 'image/svg+xml',
    }
    base_str = str(base_dir.resolve())

    # 1) <link rel="stylesheet" href="..."> → <style>...</style>
    def inline_css(m: re.Match) -> str:
        href = m.group(1)
        if href.startswith(('http', 'data', '//')):
            return m.group(0)
        css_path = (base_dir / href).resolve()
        if not str(css_path).startswith(base_str) or not css_path.exists():
            return ''
        return f'<style>{css_path.read_text(encoding="utf-8", errors="ignore")}</style>'

    html = re.sub(r'<link[^>]+href="([^"]*)"[^>]*/?\s*>', inline_css, html, flags=re.IGNORECASE)

    # 2) 기본 표 스타일 삽입 (</head> 바로 앞)
    if '</head>' in html:
        html = html.replace('</head>', _HWP_BASE_CSS + '</head>', 1)
    else:
        html = _HWP_BASE_CSS + html

    # 3) src="상대경로" → src="data:..."
    def replace_src(m: re.Match) -> str:
        src = m.group(1)
        if src.startswith(('http://', 'https://', 'data:', '//')):
            return m.group(0)
        img_path = (base_dir / src).resolve()
        if not str(img_path).startswith(base_str) or not img_path.exists():
            return m.group(0)
        ext = img_path.suffix.lstrip('.').lower()
        mime = mime_map.get(ext, 'application/octet-stream')
        data = _b64.b64encode(img_path.read_bytes()).decode()
        return f'src="data:{mime};base64,{data}"'

    return re.sub(r'src="([^"]*)"', replace_src, html)


@router.get("/files/{file_id}/hwp-preview", response_class=HTMLResponse)
async def hwp_preview(
    file_id: str,
    current_user: UserPublic = Depends(_get_user_any_auth),
):
    """HWP 파일을 HTML로 변환해 반환 (hwp5html, 이미지 base64 임베드)."""
    db = _db()
    f = await db["document_files"].find_one({"_id": parse_oid(file_id)})
    if not f:
        raise HTTPException(status_code=404, detail="파일을 찾을 수 없습니다.")

    file_path = Path(f["file_path"])
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="파일이 서버에 없습니다.")

    out_dir = tempfile.mkdtemp()
    try:
        subprocess.run(
            ["hwp5html", "--output", out_dir, str(file_path)],
            capture_output=True, text=True, timeout=60,
        )
        out_path = Path(out_dir)
        for fname in ("index.xhtml", "body.xhtml", "index.html"):
            out_file = out_path / fname
            if out_file.exists():
                html = out_file.read_text(encoding="utf-8", errors="ignore")
                html = _make_self_contained(html, out_path)
                return HTMLResponse(content=html)
        # fallback: 저장된 text_content 표시
        raise RuntimeError("hwp5html output not found")
    except Exception as e:
        logger.warning("hwp5html failed: %s", e)
        text = f.get("text_content", "변환 실패")
        escaped = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        return HTMLResponse(
            content=f"<html><body><pre style='font-family:sans-serif;line-height:1.8;padding:16px'>{escaped}</pre></body></html>"
        )
    finally:
        shutil.rmtree(out_dir, ignore_errors=True)


@router.post("/files/{file_id}/replace")
async def replace_file(
    file_id: str,
    file: UploadFile = File(...),
    current_user: UserPublic = Depends(get_current_user),
):
    """기존 파일을 새 파일로 교체 (ID/폴더 유지, 내용·텍스트 갱신)."""
    db = _db()
    f = await db["document_files"].find_one({"_id": parse_oid(file_id)})
    if not f:
        raise HTTPException(status_code=404, detail="파일을 찾을 수 없습니다.")

    content = await file.read()
    filename = file.filename or f["name"]
    extension = Path(filename).suffix

    # 기존 파일 삭제
    old_path = Path(f["file_path"])
    if old_path.exists():
        old_path.unlink()

    # 새 파일 저장
    file_uuid = str(uuid.uuid4())
    save_path = UPLOAD_DIR / f"{file_uuid}{extension}"
    save_path.write_bytes(content)

    text_content = _extract_text(content, extension)
    mime_type = mimetypes.guess_type(filename)[0] or "application/octet-stream"

    await db["document_files"].update_one(
        {"_id": parse_oid(file_id)},
        {"$set": {
            "name": filename,
            "extension": extension.lstrip(".").lower(),
            "mime_type": mime_type,
            "file_path": str(save_path),
            "text_content": text_content,
            "size": len(content),
            "updated_at": _now(),
            "updated_by": current_user.email,
        }},
    )
    updated = await db["document_files"].find_one({"_id": parse_oid(file_id)})
    return _file_out(updated)


@router.patch("/files/{file_id}")
async def update_file(
    file_id: str,
    body: dict,
    current_user: UserPublic = Depends(get_current_user),
):
    db = _db()
    f = await db["document_files"].find_one({"_id": parse_oid(file_id)})
    if not f:
        raise HTTPException(status_code=404, detail="파일을 찾을 수 없습니다.")

    update: dict = {}
    if "name" in body and body["name"]:
        update["name"] = body["name"]
    if "folder_id" in body:
        update["folder_id"] = body["folder_id"]  # None 허용 (루트로 이동)

    if not update:
        raise HTTPException(status_code=400, detail="변경할 내용이 없습니다.")

    await db["document_files"].update_one(
        {"_id": parse_oid(file_id)},
        {"$set": update},
    )
    updated = await db["document_files"].find_one({"_id": parse_oid(file_id)})
    return _file_out(updated)


@router.delete("/files/{file_id}")
async def delete_file(
    file_id: str,
    current_user: UserPublic = Depends(get_current_user),
):
    if not current_user.is_admin:
        raise HTTPException(status_code=403, detail="관리자만 삭제할 수 있습니다.")
    db = _db()
    f = await db["document_files"].find_one({"_id": parse_oid(file_id)})
    if not f:
        raise HTTPException(status_code=404, detail="파일을 찾을 수 없습니다.")

    file_path = Path(f["file_path"])
    if file_path.exists():
        file_path.unlink()

    await db["document_files"].update_one(
        {"_id": parse_oid(file_id)},
        {"$set": {"is_deleted": True}},
    )
    return {"ok": True}


@router.patch("/folders/{folder_id}")
async def update_folder(
    folder_id: str,
    name: str = Form(...),
    current_user: UserPublic = Depends(get_current_user),
):
    db = _db()
    result = await db["document_folders"].update_one(
        {"_id": parse_oid(folder_id)},
        {"$set": {"name": name.strip()}},
    )
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="폴더를 찾을 수 없습니다.")
    f = await db["document_folders"].find_one({"_id": parse_oid(folder_id)})
    return {
        "id": str(f["_id"]),
        "name": f["name"],
        "parent_id": f.get("parent_id"),
        "created_at": fmt_dt(f.get("created_at")),
    }


@router.post("/files/{file_id}/convert-to-docx")
async def convert_to_docx(
    file_id: str,
    current_user: UserPublic = Depends(get_current_user),
):
    """HWP 파일을 DOCX로 변환하여 같은 폴더에 새 파일로 저장합니다."""
    db = _db()
    f = await db["document_files"].find_one({"_id": parse_oid(file_id)})
    if not f:
        raise HTTPException(status_code=404, detail="파일을 찾을 수 없습니다.")
    if f.get("extension", "").lower() != "hwp":
        raise HTTPException(status_code=400, detail="HWP 파일만 변환할 수 있습니다.")

    src = Path(f["file_path"])
    with tempfile.TemporaryDirectory() as tmpdir:
        out = _lo_convert(src, "docx", Path(tmpdir))
        content = out.read_bytes()

    new_name = Path(f["name"]).stem + ".docx"
    file_uuid = str(uuid.uuid4())
    dest = UPLOAD_DIR / f"{file_uuid}.docx"
    dest.write_bytes(content)
    text = _extract_text(content, ".docx")

    doc = {
        "name": new_name,
        "folder_id": f.get("folder_id"),
        "extension": "docx",
        "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "file_path": str(dest),
        "text_content": text,
        "size": len(content),
        "is_deleted": False,
        "converted_from": "hwp",
        "created_at": _now(),
        "created_by": current_user.email,
    }
    result = await db["document_files"].insert_one(doc)
    doc["_id"] = result.inserted_id
    return _file_out(doc)


@router.get("/files/{file_id}/edit-content")
async def get_edit_content(
    file_id: str,
    current_user: UserPublic = Depends(get_current_user),
):
    """편집용 콘텐츠 반환 — txt/md/csv: text, docx: HTML"""
    db = _db()
    f = await db["document_files"].find_one({"_id": parse_oid(file_id)})
    if not f:
        raise HTTPException(status_code=404, detail="파일을 찾을 수 없습니다.")

    ext = f.get("extension", "").lower()
    src = Path(f["file_path"])

    if ext in ("txt", "md", "csv"):
        text = src.read_text(encoding="utf-8", errors="replace")
        return {"content_type": "text", "content": text}
    elif ext == "docx":
        import mammoth
        with open(src, "rb") as fp:
            result = mammoth.convert_to_html(fp)
        return {"content_type": "html", "content": result.value}
    else:
        raise HTTPException(status_code=400, detail="편집을 지원하지 않는 파일 형식입니다.")


@router.put("/files/{file_id}/edit-content")
async def save_edit_content(
    file_id: str,
    content_type: str = Form(...),
    content: str = Form(...),
    current_user: UserPublic = Depends(get_current_user),
):
    """편집 내용 저장 — text: 직접 저장, html: LibreOffice로 DOCX 변환 후 저장"""
    db = _db()
    f = await db["document_files"].find_one({"_id": parse_oid(file_id)})
    if not f:
        raise HTTPException(status_code=404, detail="파일을 찾을 수 없습니다.")

    file_path = Path(f["file_path"])

    if content_type == "text":
        file_path.write_text(content, encoding="utf-8")
        new_text = content
    elif content_type == "html":
        with tempfile.TemporaryDirectory() as tmpdir:
            html_file = Path(tmpdir) / "content.html"
            html_file.write_text(f"<html><head><meta charset='utf-8'></head><body>{content}</body></html>", encoding="utf-8")
            docx_file = _lo_convert(html_file, "docx", Path(tmpdir))
            file_path.write_bytes(docx_file.read_bytes())
        new_text = _extract_text(file_path.read_bytes(), ".docx")
    else:
        raise HTTPException(status_code=400, detail="알 수 없는 content_type입니다.")

    await db["document_files"].update_one(
        {"_id": parse_oid(file_id)},
        {"$set": {"text_content": new_text, "size": file_path.stat().st_size}},
    )
    updated = await db["document_files"].find_one({"_id": parse_oid(file_id)})
    return _file_out(updated)


@router.delete("/folders/{folder_id}")
async def delete_folder(
    folder_id: str,
    current_user: UserPublic = Depends(get_current_user),
):
    if not current_user.is_admin:
        raise HTTPException(status_code=403, detail="관리자만 삭제할 수 있습니다.")
    db = _db()
    count = await db["document_files"].count_documents(
        {"folder_id": folder_id, "is_deleted": {"$ne": True}}
    )
    if count > 0:
        raise HTTPException(
            status_code=400, detail="폴더 안에 파일이 있습니다. 파일을 먼저 삭제해주세요."
        )
    await db["document_folders"].delete_one({"_id": parse_oid(folder_id)})
    return {"ok": True}
